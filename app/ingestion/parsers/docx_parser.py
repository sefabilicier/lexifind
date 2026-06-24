"""
DOCX parser using python-docx.
Extracts paragraphs and table content with heading awareness.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.ingestion.parsers.base import BaseParser, ParsedDocument
from app.observability.logger import get_logger

logger = get_logger(__name__)


class DOCXParser(BaseParser):

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: Path) -> ParsedDocument:
        logger.info(
            "Parsing DOCX",
            file=file_path.name,
        )

        doc = Document(str(file_path))
        blocks = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Mark headings for structural awareness
            if para.style.name.startswith("Heading"):
                blocks.append(f"\n## {text}\n")
            else:
                blocks.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    blocks.append(row_text)

        full_content = "\n".join(blocks)

        metadata = {
            "source": file_path.name,
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        logger.info(
            "DOCX parsed successfully",
            file=file_path.name,
            chars=len(full_content),
        )

        return ParsedDocument(
            content=full_content,
            metadata=metadata,
            pages=[full_content],
        )